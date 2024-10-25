import streamlit as st  
import requests  
import PyPDF2  
from azure.ai.formrecognizer import DocumentAnalysisClient  
from azure.core.credentials import AzureKeyCredential  
from openai import AzureOpenAI  
from docx import Document  
import io  
from surya.ocr import run_ocr  
from surya.model.detection.model import load_model as load_det_model  
from surya.model.detection.model import load_processor as load_det_processor  
from surya.model.recognition.model import load_model as load_rec_model  
from surya.model.recognition.processor import load_processor as load_rec_processor  
import fitz  # PyMuPDF  
from PIL import Image  
summarizer_endpoint = "https://theswedes.openai.azure.com/openai/deployments/GPT-4-Omni/chat/completions?api-version=2024-02-15-preview"  
summarizer_api_key = "783973291a7c4a74a1120133309860c0"  
  
summarizer_client = AzureOpenAI(  
    azure_endpoint=summarizer_endpoint,  
    api_key=summarizer_api_key,  
    api_version="2024-02-15-preview",  
)  
  
# Function to extract insights using Phi-3.5-Vision  
def extract_insights_phi_vision(pdf_data):  
    endpoint = "https://Phi-3-medium-128k-instruct-xadpw.eastus2.models.ai.azure.com/v1/chat/completions"  
    headers = {  
        'Content-Type': 'application/json',  
        'Authorization': 'Bearer lzWrm9dPkXatZ7JKoC8nIf4Jjn1fjSRQ'  
    }  
  
    insights = []  
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_data))  
    for page_num, page in enumerate(pdf_reader.pages):  
        page_text = page.extract_text()  
        page_images = []  # Placeholder for image extraction logic  
        data = {  
            "messages": [  
                {"role": "system", "content": "Extract essential text and image insights:"},  
                {"role": "user", "content": f"Text: {page_text}"},  
                {"role": "user", "content": f"Images: {page_images}"}  
            ]  
        }  
        try:  
            response = requests.post(endpoint, headers=headers, json=data)  
            response.raise_for_status()  
            content = response.json()['choices'][0]['message']['content']  
            insights.append((page_num + 1, content))  # Store page number and content  
        except requests.exceptions.RequestException as e:  
            st.error(f"Error on page {page_num + 1}: {str(e)}")  
    return insights  
  

def extract_text_from_pdf(pdf_data):  
    form_recognizer_endpoint = "https://patentocr.cognitiveservices.azure.com/"  
    form_recognizer_api_key = "cd6b8996d93447be88d995729c924bcb"  
  
    try:  
        document_analysis_client = DocumentAnalysisClient(  
            endpoint=form_recognizer_endpoint,  
            credential=AzureKeyCredential(form_recognizer_api_key),  
        )  
        poller = document_analysis_client.begin_analyze_document(  
            "prebuilt-document", document=pdf_data  
        )  
        result = poller.result()  
  
        text_by_page = {}  
        for page in result.pages:  
            page_text = ""  
            for line in page.lines:  
                page_text += line.content + "\n"  
            text_by_page[page.page_number] = page_text  
  
        return text_by_page  
    except Exception as e:  
        st.error(f"An error occurred: {str(e)}")  
        return None  
  
# New function to extract text using Surya OCR  
def extract_text_with_surya(pdf_data):  
    text_by_page = {}  
    det_processor, det_model = load_det_processor(), load_det_model()  
    rec_model, rec_processor = load_rec_model(), load_rec_processor()  
  
    # Open the PDF file with PyMuPDF  
    doc = fitz.open(stream=pdf_data, filetype="pdf")  
  
    for page_num in range(len(doc)):  
        try:  
            page = doc.load_page(page_num)  
            # Render page to an image  
            pix = page.get_pixmap()  
            image = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)  
  
            # Use Surya OCR to extract text and image insights  
            predictions = run_ocr([image], [["en"]], det_model, det_processor, rec_model, rec_processor)  
  
            # Debugging: print the structure of OCRResult  
            st.write(predictions)  # Use st.write to output the structure  
  
            # Adjust the following line based on the actual structure of predictions  
            for result in predictions:  
                if hasattr(result, 'text_lines'):  
                    page_text = "\n".join([line.text for line in result.text_lines])  
                else:  
                    # If the structure is different, adjust accordingly  
                    page_text = "\n".join([line['text'] for line in result['text_lines']])  
                  
                text_by_page[page_num + 1] = page_text  
        except Exception as e:  
            st.error(f"Error processing page {page_num + 1} with Surya: {str(e)}")  
  
    return text_by_page  
  
# Function to compare insights using Azure OpenAI  
def compare_insights(insights_phi, insights_azure, insights_surya):  
    comparison_prompt = (  
        "Compare the following insights:\n"  
        "\nText Extraction:\n"  
        "Evaluate the accuracy, clarity, and completeness of text extracted by each approach.\n"  
        "Phi-3.5-Vision Insights:\n"  
        f"{insights_phi}\n"  
        "Azure Document Intelligence Insights:\n"  
        f"{insights_azure}\n"  
        "Surya OCR Insights:\n"  
        f"{insights_surya}\n"  
        "\nImage and Visual Elements Detailing:\n"  
        "Evaluate how well each method extracts content from images or visual elements.\n"  
        "Provide a detailed analysis of which approach delivers superior results for text and image content."  
    )  
  
    try:  
        completion = summarizer_client.chat.completions.create(  
            model="GPT-4-Omni",  
            messages=[  
                {"role": "system", "content": "You are a helpful assistant."},  
                {"role": "user", "content": comparison_prompt}  
            ],  
            max_tokens=4096,  
            temperature=0,  
            top_p=1,  
            frequency_penalty=0,  
            presence_penalty=0,  
            stop=None,  
            stream=False  
        )  
        return completion.choices[0].message.content  
    except Exception as e:  
        st.error(f"Error comparing insights: {str(e)}")  
        return "Comparison failed."  
  
# Function to create a Word document  
def create_word_document(insights_phi, insights_azure, insights_surya, comparison_result):  
    doc = Document()  
    doc.add_heading('PDF Insight Comparison', 0)  
  
    # Add insights from Phi-3.5-Vision with page numbers  
    doc.add_heading('Insights from Phi-3.5-Vision', level=1)  
    for page_num, content in insights_phi:  
        doc.add_heading(f'Page {page_num}', level=2)  
        doc.add_paragraph(content)  
  
    # Add insights from Azure Document Intelligence with page numbers  
    doc.add_heading('Insights from Azure Document Intelligence', level=1)  
    for page_num, content in insights_azure.items():  
        doc.add_heading(f'Page {page_num}', level=2)  
        doc.add_paragraph(content)  
  
    # Add insights from Surya OCR with page numbers  
    doc.add_heading('Insights from Surya OCR', level=1)  
    for page_num, content in insights_surya.items():  
        doc.add_heading(f'Page {page_num}', level=2)  
        doc.add_paragraph(content)  
  
    # Add comparison results  
    doc.add_heading('Comparison Results', level=1)  
    doc.add_paragraph(comparison_result)  
  
    # Add a table for comparison results  
    table = doc.add_table(rows=1, cols=3)  
    hdr_cells = table.rows[0].cells  
    hdr_cells[0].text = 'Source'  
    hdr_cells[1].text = 'Page Number'  
    hdr_cells[2].text = 'Content'  
  
    for page_num, content in insights_phi:  
        row_cells = table.add_row().cells  
        row_cells[0].text = 'Phi-3.5-Vision'  
        row_cells[1].text = str(page_num)  
        row_cells[2].text = content  
  
    for page_num, content in insights_azure.items():  
        row_cells = table.add_row().cells  
        row_cells[0].text = 'Azure Document Intelligence'  
        row_cells[1].text = str(page_num)  
        row_cells[2].text = content  
  
    for page_num, content in insights_surya.items():  
        row_cells = table.add_row().cells  
        row_cells[0].text = 'Surya OCR'  
        row_cells[1].text = str(page_num)  
        row_cells[2].text = content  
  
    return doc  
  
st.title("PDF Insight Comparison App")  
uploaded_file = st.file_uploader("Upload a PDF file", type="pdf")  
  
if uploaded_file is not None:  
    pdf_data = uploaded_file.read()  
  
    with st.spinner("Extracting insights using Phi-3.5-Vision..."):  
        insights_phi = extract_insights_phi_vision(pdf_data)  
  
    with st.spinner("Extracting insights using Azure Document Intelligence..."):  
        insights_azure_text = extract_text_from_pdf(pdf_data)  
  
    with st.spinner("Extracting insights using Surya OCR..."):  
        insights_surya_text = extract_text_with_surya(pdf_data)  
  
    if insights_azure_text and insights_surya_text:  
        with st.spinner("Comparing insights..."):  
            comparison_result = compare_insights(insights_phi, insights_azure_text, insights_surya_text)  
  
        st.write("### Comparison Results")  
        st.write(comparison_result)  
  
        # Create and download Word document  
        doc = create_word_document(insights_phi, insights_azure_text, insights_surya_text, comparison_result)  
        doc_file = "comparison_results.docx"  
        doc.save(doc_file)  
  
        with open(doc_file, "rb") as file:  
            st.download_button(  
                label="Download Comparison Results as Word Document",  
                data=file,  
                file_name=doc_file,  
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"  
            )
